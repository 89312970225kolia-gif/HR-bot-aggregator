from __future__ import annotations

import asyncio
import re
from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from pypdf import PdfReader


class ResumeParseError(ValueError):
    """The supplied resume cannot be parsed into non-empty text."""


def normalize_text(text: str) -> str:
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line).strip()


def _extract_docx(file_bytes: bytes) -> str:
    try:
        document = Document(BytesIO(file_bytes))
    except (PackageNotFoundError, BadZipFile, ValueError, KeyError, OSError) as error:
        raise ResumeParseError("DOCX file is damaged or invalid") from error

    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return normalize_text("\n".join(parts))


def _extract_pdf(file_bytes: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(file_bytes))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as error:
        raise ResumeParseError("PDF file is damaged or invalid") from error
    return normalize_text(text)


async def extract_resume_text(
    file_bytes: bytes, filename: str, mime_type: str | None
) -> str:
    suffix = Path(filename).suffix.lower()
    normalized_mime = (mime_type or "").lower()
    if suffix == ".docx" or normalized_mime.endswith("wordprocessingml.document"):
        result = await asyncio.to_thread(_extract_docx, file_bytes)
    elif suffix == ".pdf" or normalized_mime == "application/pdf":
        result = await asyncio.to_thread(_extract_pdf, file_bytes)
    else:
        raise ResumeParseError("Unsupported resume format")
    if not result:
        if suffix == ".pdf":
            raise ResumeParseError(
                "Не удалось извлечь текст из PDF. Попробуйте отправить текстовое PDF или DOCX."
            )
        raise ResumeParseError("DOCX contains no extractable text")
    return result
