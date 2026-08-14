from io import BytesIO

import pytest
from docx import Document
from reportlab.pdfgen import canvas

from app.services.resume_parser import ResumeParseError, extract_resume_text


@pytest.mark.asyncio
async def test_docx_paragraphs_and_tables_are_extracted() -> None:
    document = Document()
    document.add_paragraph("Опыт с ChatGPT и Kling")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Портфолио"
    table.cell(0, 1).text = "example.com"
    buffer = BytesIO()
    document.save(buffer)

    text = await extract_resume_text(
        buffer.getvalue(),
        "resume.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert "Опыт с ChatGPT" in text
    assert "Портфолио | example.com" in text


@pytest.mark.asyncio
async def test_pdf_text_is_extracted() -> None:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer)
    pdf.drawString(72, 750, "AI content portfolio")
    pdf.save()

    text = await extract_resume_text(buffer.getvalue(), "resume.pdf", "application/pdf")

    assert "AI content portfolio" in text


@pytest.mark.asyncio
async def test_scanned_or_blank_pdf_has_controlled_error() -> None:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer)
    pdf.showPage()
    pdf.save()

    with pytest.raises(ResumeParseError, match="Не удалось извлечь текст"):
        await extract_resume_text(buffer.getvalue(), "scan.pdf", "application/pdf")


@pytest.mark.asyncio
async def test_broken_docx_has_controlled_error() -> None:
    with pytest.raises(ResumeParseError, match="damaged|invalid"):
        await extract_resume_text(
            b"not-a-docx",
            "resume.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
